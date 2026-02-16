"""Word document generation service using python-docx"""

import io
import re

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def html_to_docx(html_content: str, title: str = "Document") -> bytes:
    """Convert HTML content to DOCX bytes

    Simple HTML-to-DOCX converter that handles basic tags.

    Args:
        html_content: HTML string to convert
        title: Document title

    Returns:
        DOCX file as bytes
    """
    doc = DocxDocument()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'MS Gothic'
    font.size = Pt(10.5)

    # Parse HTML and convert to docx elements
    # Simple regex-based parser for basic HTML
    text = _strip_tags_to_structured(html_content)

    for line in text:
        if line["type"] == "h1":
            p = doc.add_heading(line["text"], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line["type"] == "h2":
            doc.add_heading(line["text"], level=2)
        elif line["type"] == "h3":
            doc.add_heading(line["text"], level=3)
        elif line["type"] == "text":
            if line["text"].strip():
                doc.add_paragraph(line["text"])
        elif line["type"] == "table_row":
            # For simple key-value pairs, add as paragraph
            doc.add_paragraph(line["text"])

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.read()


def _strip_tags_to_structured(html: str) -> list[dict]:
    """Parse HTML into structured text elements"""
    result = []

    # Remove script/style tags
    html = re.sub(r'<(script|style)[^>]*>.*?</\1>', '', html, flags=re.DOTALL | re.IGNORECASE)

    # Process headings
    for tag, level in [('h1', 'h1'), ('h2', 'h2'), ('h3', 'h3')]:
        html = re.sub(
            rf'<{tag}[^>]*>(.*?)</{tag}>',
            lambda m, l=level: f'\n[{l}]{_clean_text(m.group(1))}[/{l}]\n',
            html,
            flags=re.DOTALL | re.IGNORECASE,
        )

    # Process table rows
    html = re.sub(
        r'<tr[^>]*>(.*?)</tr>',
        lambda m: '\n[row]' + _clean_text(re.sub(r'<t[hd][^>]*>(.*?)</t[hd]>', r' \1 |', m.group(1), flags=re.DOTALL | re.IGNORECASE)) + '[/row]\n',
        html,
        flags=re.DOTALL | re.IGNORECASE,
    )

    # Process line breaks
    html = re.sub(r'<br\s*/?\s*>', '\n', html, flags=re.IGNORECASE)
    html = re.sub(r'</p>', '\n', html, flags=re.IGNORECASE)
    html = re.sub(r'</div>', '\n', html, flags=re.IGNORECASE)
    html = re.sub(r'</li>', '\n', html, flags=re.IGNORECASE)

    # Remove remaining tags
    html = re.sub(r'<[^>]+>', '', html)

    # Decode entities
    html = html.replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')

    for line in html.split('\n'):
        line = line.strip()
        if not line:
            continue

        if line.startswith('[h1]') and line.endswith('[/h1]'):
            result.append({"type": "h1", "text": line[4:-5].strip()})
        elif line.startswith('[h2]') and line.endswith('[/h2]'):
            result.append({"type": "h2", "text": line[4:-5].strip()})
        elif line.startswith('[h3]') and line.endswith('[/h3]'):
            result.append({"type": "h3", "text": line[4:-5].strip()})
        elif line.startswith('[row]') and line.endswith('[/row]'):
            result.append({"type": "table_row", "text": line[5:-6].strip().rstrip('|')})
        else:
            result.append({"type": "text", "text": line})

    return result


def _clean_text(text: str) -> str:
    """Remove HTML tags from text"""
    return re.sub(r'<[^>]+>', '', text).strip()
